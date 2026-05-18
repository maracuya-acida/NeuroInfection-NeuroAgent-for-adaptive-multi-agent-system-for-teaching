"""
- Educational Agent with Ollama  
Flask backend for generating lesson plans, syllabi, and curriculum maps
"""

from flask import Flask, request, jsonify, send_file, render_template, send_from_directory
from flask_cors import CORS
import requests
import re
from datetime import datetime
import os
import json

# For exporting to DOCX and PDF
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# ─────────────────────────────────────────────
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "gemma:2b" 
OUTPUT_DIR = "outputs"

# Create necessary folders if they don't exist
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Flask configuration to find your index.html
app = Flask(__name__, template_folder='.') 
CORS(app)


# ─────────────────────────────────────────────
# KNOWLEDGE LOADING FROM JSON
# (based on embeddingforjson.py, without external embeddings)
# ─────────────────────────────────────────────

def load_knowledge_base(json_path: str = "relationaldata.json") -> str:
    """
    Reads relationaldata.json and builds a rich text block
    with all topics, definitions, formulas, examples, and methodologies.
    Returns the context as a string ready to inject into the system prompt.
    """
    if not os.path.exists(json_path):
        print(f"⚠️ {json_path} not found. The agent will operate without a knowledge base.")
        return ""

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    sections = []

    # ── Calculus Topics ──────────────────────────────────────────────────
    calculus = data.get("calculus", {})
    topics   = calculus.get("topics", {})

    for topic_id, topic in topics.items():
        block = f"""
=== {topic.get('display_name', topic_id)} ===
Topic ID: {topic_id}

Definition:
{topic.get('definition', '')}

Formulas:
{json.dumps(topic.get('formulas', {}), indent=2, ensure_ascii=False)}

Examples:
{json.dumps(topic.get('examples', {}), indent=2, ensure_ascii=False)}
""".strip()
        sections.append(block)

    # ── Pedagogical Methodology ────────────────────────────────────────────
    pedagogy = data.get("pedagogical_methodology", {})
    if pedagogy:
        methods_block = "=== Pedagogical Methodology ===\n"

        for method_key, method in pedagogy.get("methods", {}).items():
            methods_block += f"\n[{method.get('label', method_key)}]\n"
            methods_block += f"{method.get('description', '')}\n"
            if "principles" in method:
                methods_block += "Principles:\n" + "\n".join(f"  - {p}" for p in method["principles"]) + "\n"
            if "target_profile" in method:
                methods_block += "Target profile: " + ", ".join(method["target_profile"]) + "\n"

        # UDL Applications
        udl_apps = pedagogy.get("udl_application", {})
        for app_key, udl in udl_apps.items():
            methods_block += f"\n[UDL applied: {udl.get('label', app_key)}]\n"
            methods_block += "Representation:\n" + "\n".join(f"  - {r}" for r in udl.get("representation", [])) + "\n"
            methods_block += "Action and expression:\n" + "\n".join(f"  - {a}" for a in udl.get("action_and_expression", [])) + "\n"
            methods_block += "Engagement:\n" + "\n".join(f"  - {e}" for e in udl.get("engagement", [])) + "\n"

        sections.append(methods_block.strip())

    knowledge_text = "\n\n" + ("─" * 60) + "\n\n".join(sections)
    print(f"✅ Knowledge base loaded: {len(topics)} Calculus topics + pedagogical methodology.")
    return knowledge_text


# Single load when the server starts
KNOWLEDGE_BASE: str = load_knowledge_base("relationaldata.json")

# Keywords that trigger JSON context injection
KNOWLEDGE_KEYWORDS = [
    "calculus", "integral", "derivative", "limit", "limit",
    "trigonometric", "trigonometry", "substitution", "parts",
    "antiderivative", "fundamental theorem", "syllabus", "curriculum",
    "pedagogy", "udl", "cra", "neurodivergent", "adaptive"
]

def build_system_prompt(topic: str = "") -> str:
    """
    Builds the base system prompt. If the topic contains keywords
    related to the JSON, adds the knowledge context.
    """
    base = SYSTEM_BASE

    if not KNOWLEDGE_BASE:
        return base

    topic_lower = (topic or "").lower()
    should_inject = any(kw in topic_lower for kw in KNOWLEDGE_KEYWORDS)

    if should_inject:
        context_block = f"""

═══════════════════════════════════════════
AVAILABLE KNOWLEDGE BASE (use as technical reference):
═══════════════════════════════════════════
{KNOWLEDGE_BASE}
═══════════════════════════════════════════
"""
        return base + context_block

    return base


# ─────────────────────────────────────────────
# OLLAMA UTILITIES
# ─────────────────────────────────────────────

def query_ollama(prompt: str, system_prompt: str = "") -> str:
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "system": system_prompt,
        "stream": False,
        "context": [],  # <--- ADD THIS LINE to clear the AI's memory
        "options": {
            "temperature": 0.4,
            "num_thread": 4, # Use real threads from your PC
            "num_ctx": 4096  # Limits the working memory size
        }
    }
    try:
        resp = requests.post(OLLAMA_URL, json=payload, timeout=600)
        resp.raise_for_status()
        data = resp.json()
        return data.get("response", "").strip()
    except requests.exceptions.ConnectionError:
        return "ERROR_OLLAMA_OFFLINE"
    except Exception as e:
        return f"ERROR: {str(e)}"


# ─────────────────────────────────────────────
# AGENT PROMPTS
# ─────────────────────────────────────────────

SYSTEM_BASE = """You are NeuroAgent. Generate expert pedagogical content.
CRITICAL RULE: Respond directly with Markdown format.
Do not greet, do not explain what you are going to do, do not give introductions.
Be brief, technical, and direct in each section."""


def prompt_lesson_plan(topic: str, approach: str, learning: str, level: str, duration: str) -> str:
    return f"""Act as an expert Instructional Designer. Your task is to create a HIGH-IMPACT Lesson Plan.

KEY DATA:
- TOPIC: {topic}
- APPROACH: {approach}
- OBJECTIVE: {learning}
- LEVEL: {level}
- TIME: {duration}

FORMATTING INSTRUCTIONS:
1. Use Markdown with headings (##) and subheadings (###).
2. The didactic sequence MUST sum exactly {duration}.
3. Divide activities into: Beginning (Motivation), Development (Construction), and Closing (Metacognition).

REQUIRED CONTENT:
## 1. Identification
## 2. Objectives (General and 3 specific)
## 3. Content (Knowledge, Know-how, Being)
## 4. Didactic Sequence (Detail minute by minute)
## 5. Assessment (Define a tangible product the student will deliver)
## 6. Required Resources

Be creative, innovative, and ensure activities are coherent with the {level} level."""


def prompt_syllabus(topic: str, approach: str, learning: str, level: str, weeks: int) -> str:
    return f"""Generate a complete academic Syllabus with the following data:

- **Subject**: {topic}
- **Approach**: {approach}
- **Core competency**: {learning}
- **Level**: {level}
- **Duration**: {weeks} weeks

The syllabus must include:
1. Identification data (subject, code, credits, hours)
2. General course description
3. Competencies and learning outcomes
4. Thematic units with weekly content
5. Teaching-learning methodology
6. Assessment system (exams, assignments, final exam with percentages)
7. Bibliography (basic and complementary)
8. Detailed weekly schedule

Use structured markdown format. Be precise with percentages and dates."""


def prompt_curriculum(topic: str, approach: str, learning: str, level: str, cycles: int) -> str:
    return f"""Generate a complete Curriculum Map with the following data:

- **Career/Program**: {topic}
- **Educational approach**: {approach}
- **Graduate profile**: {learning}
- **Level**: {level}
- **Cycles/Semesters**: {cycles}

The curriculum map must include:
1. Program presentation
2. Entry and exit profile
3. Program objectives
4. Curriculum structure by cycles (with subjects, credits, and hours)
5. Training areas (basic, professional, specialization, electives)
6. Main prerequisite map
7. Total workload
8. Professional profile and job market

Organize subjects in a table per cycle. Include credits for each subject."""


def prompt_adaptive(topic: str, approach: str, learning: str, level: str, duration: str, style: str) -> str:
    return f"""Generate a personalized ADAPTIVE Lesson Plan:

- **Topic**: {topic}
- **Approach**: {approach}  
- **Expected learning**: {learning}
- **Level**: {level}
- **Duration**: {duration}
- **Predominant learning style**: {style}

Adapt the plan considering the indicated learning style:
- If Visual: include described diagrams, concept maps, graphic organizers
- If Auditory: include debates, presentations, podcasts, group discussion
- If Kinesthetic: include practical activities, experiments, role-play
- If Reading/Writing: include text analysis, essays, research

Structure the same as a standard lesson plan but with specific activities for that style.
Include an "Adaptations for other styles" section at the end."""


# ─────────────────────────────────────────────
# EXPORTERS
# ─────────────────────────────────────────────

def markdown_to_plain(md_text: str) -> list[dict]:
    """Converts markdown to a list of {type, text} blocks."""
    blocks = []
    for line in md_text.split('\n'):
        line = line.rstrip()
        if line.startswith('### '):
            blocks.append({'type': 'h3', 'text': line[4:]})
        elif line.startswith('## '):
            blocks.append({'type': 'h2', 'text': line[3:]})
        elif line.startswith('# '):
            blocks.append({'type': 'h1', 'text': line[2:]})
        elif line.startswith('- ') or line.startswith('* '):
            blocks.append({'type': 'bullet', 'text': line[2:]})
        elif re.match(r'^\d+\. ', line):
            blocks.append({'type': 'numbered', 'text': re.sub(r'^\d+\. ', '', line)})
        elif line.startswith('**') and line.endswith('**'):
            blocks.append({'type': 'bold', 'text': line[2:-2]})
        elif line == '' or line == '---':
            blocks.append({'type': 'space', 'text': ''})
        else:
            # Clean inline markdown
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            clean = re.sub(r'`(.*?)`', r'\1', clean)
            if clean.strip():
                blocks.append({'type': 'para', 'text': clean})
    return blocks


def export_to_docx(content: str, title: str, filename: str) -> str:
    """Exports markdown content to DOCX."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")

    doc = Document()

    # Basic styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Document title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    blocks = markdown_to_plain(content)
    for block in blocks:
        t = block['text'].strip()
        if not t and block['type'] == 'space':
            doc.add_paragraph()
            continue
        if not t:
            continue

        if block['type'] == 'h1':
            doc.add_heading(t, level=1)
        elif block['type'] == 'h2':
            doc.add_heading(t, level=2)
        elif block['type'] == 'h3':
            doc.add_heading(t, level=3)
        elif block['type'] == 'bullet':
            p = doc.add_paragraph(t, style='List Bullet')
        elif block['type'] == 'numbered':
            p = doc.add_paragraph(t, style='List Number')
        elif block['type'] == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(t)
            run.bold = True
        else:
            doc.add_paragraph(t)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


def export_to_pdf(content: str, title: str, filename: str) -> str:
    """Exports markdown content to PDF with ReportLab."""
    if not PDF_AVAILABLE:
        raise ImportError("reportlab not installed")

    path = os.path.join(OUTPUT_DIR, filename)
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=60, leftMargin=60,
                            topMargin=60, bottomMargin=60)

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle('CustomTitle',
        parent=styles['Title'],
        fontSize=20, spaceAfter=6,
        textColor=colors.HexColor('#1a1a2e'),
        alignment=TA_CENTER)

    date_style = ParagraphStyle('DateStyle',
        parent=styles['Normal'],
        fontSize=9, spaceAfter=20,
        textColor=colors.HexColor('#888888'),
        alignment=TA_CENTER)

    h1_style = ParagraphStyle('H1',
        parent=styles['Heading1'],
        fontSize=16, spaceBefore=16, spaceAfter=6,
        textColor=colors.HexColor('#1a1a2e'))

    h2_style = ParagraphStyle('H2',
        parent=styles['Heading2'],
        fontSize=13, spaceBefore=12, spaceAfter=4,
        textColor=colors.HexColor('#16213e'))

    h3_style = ParagraphStyle('H3',
        parent=styles['Heading3'],
        fontSize=11, spaceBefore=8, spaceAfter=4,
        textColor=colors.HexColor('#0f3460'))

    body_style = ParagraphStyle('Body',
        parent=styles['Normal'],
        fontSize=10, spaceAfter=4,
        leading=14, alignment=TA_JUSTIFY)

    bullet_style = ParagraphStyle('Bullet',
        parent=styles['Normal'],
        fontSize=10, spaceAfter=3,
        leftIndent=20, leading=13)

    # Title
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}", date_style))

    blocks = markdown_to_plain(content)
    for block in blocks:
        t = block['text'].strip()
        if not t and block['type'] == 'space':
            story.append(Spacer(1, 6))
            continue
        if not t:
            continue

        # Escape special characters for ReportLab
        t_safe = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        if block['type'] == 'h1':
            story.append(Paragraph(t_safe, h1_style))
        elif block['type'] == 'h2':
            story.append(Paragraph(t_safe, h2_style))
        elif block['type'] == 'h3':
            story.append(Paragraph(t_safe, h3_style))
        elif block['type'] in ('bullet', 'numbered'):
            prefix = '• ' if block['type'] == 'bullet' else '→ '
            story.append(Paragraph(f"{prefix}{t_safe}", bullet_style))
        elif block['type'] == 'bold':
            story.append(Paragraph(f"<b>{t_safe}</b>", body_style))
        else:
            story.append(Paragraph(t_safe, body_style))

    doc.build(story)
    return path


# ─────────────────────────────────────────────
# API ROUTES
# ─────────────────────────────────────────────
# ====================== MAIN ROUTE ======================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/<path:path>')
def serve_static(path):
    """Serves static files like index.html if necessary"""
    return send_from_directory('.', path)

@app.route('/api/greeting', methods=['GET'])
def greeting():
    name = request.args.get('name', 'Teacher')
    hour = datetime.now().hour
    if hour < 12:
        greeting_text = "Good morning"
    elif hour < 18:
        greeting_text = "Good afternoon"
    else:
        greeting_text = "Good evening"

    prompt = f"""Give a warm and professional greeting to {name}.
The greeting is: {greeting_text}.
Introduce yourself as NeuroAgent, pedagogical assistant.
Briefly mention what you can do (lesson plan, syllabus, curriculum map).
Maximum 3 sentences. Be friendly but direct."""

    response = query_ollama(prompt, SYSTEM_BASE)
    if response.startswith("ERROR"):
        response = f"{greeting_text}, {name}! I'm NeuroAgent, your pedagogical assistant. I can help you create lesson plans, syllabi, and curriculum maps."

    return jsonify({"greeting": response, "hour": greeting_text})


@app.route('/api/lesson-plan', methods=['POST'])
def lesson_plan():
    data = request.json
    topic = data.get('topic', '')
    approach = data.get('approach', '')
    learning = data.get('learning', '')
    level = data.get('level', 'High School')
    duration = data.get('duration', '60 minutes')
    adaptive = data.get('adaptive', False)
    style = data.get('style', 'Visual')

    if not topic or not approach or not learning:
        return jsonify({"error": "Missing required data"}), 400

    if adaptive:
        prompt = prompt_adaptive(topic, approach, learning, level, duration, style)
    else:
        prompt = prompt_lesson_plan(topic, approach, learning, level, duration)

    # Inject JSON context if the topic is relevant
    system = build_system_prompt(topic)
    content = query_ollama(prompt, system)

    if content.startswith("ERROR_OLLAMA_OFFLINE"):
        return jsonify({"error": "Ollama is not running. Start with: ollama serve"}), 503

    return jsonify({
        "type": "lesson_plan",
        "title": f"Lesson Plan: {topic}",
        "content": content,
        "adaptive": adaptive
    })


@app.route('/api/syllabus', methods=['POST'])
def syllabus():
    data = request.json
    topic = data.get('topic', '')
    approach = data.get('approach', '')
    learning = data.get('learning', '')
    level = data.get('level', 'University')
    weeks = data.get('weeks', 16)

    if not topic or not approach or not learning:
        return jsonify({"error": "Missing required data"}), 400

    prompt = prompt_syllabus(topic, approach, learning, level, weeks)

    # Inject JSON context if the topic is relevant
    system = build_system_prompt(topic)
    content = query_ollama(prompt, system)

    if content.startswith("ERROR_OLLAMA_OFFLINE"):
        return jsonify({"error": "Ollama is not running. Start with: ollama serve"}), 503

    return jsonify({
        "type": "syllabus",
        "title": f"Syllabus: {topic}",
        "content": content
    })


@app.route('/api/curriculum', methods=['POST'])
def curriculum():
    data = request.json
    topic = data.get('topic', '')
    approach = data.get('approach', '')
    learning = data.get('learning', '')
    level = data.get('level', 'University')
    cycles = data.get('cycles', 8)

    if not topic or not approach or not learning:
        return jsonify({"error": "Missing required data"}), 400

    prompt = prompt_curriculum(topic, approach, learning, level, cycles)

    # Inject JSON context if the topic is relevant
    system = build_system_prompt(topic)
    content = query_ollama(prompt, system)

    if content.startswith("ERROR_OLLAMA_OFFLINE"):
        return jsonify({"error": "Ollama is not running. Start with: ollama serve"}), 503

    return jsonify({
        "type": "curriculum",
        "title": f"Curriculum Map: {topic}",
        "content": content
    })


@app.route('/api/export', methods=['POST'])
def export():
    data = request.json
    format_type = data.get('format', 'pdf')  # 'pdf' or 'docx'
    title = data.get('title', 'Educational Document')
    content = data.get('content', '')
    doc_type = data.get('type', 'document')

    if not content:
        return jsonify({"error": "No content to export"}), 400

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:40]
    filename = f"{safe_title}_{timestamp}"

    try:
        if format_type == 'docx':
            filename += '.docx'
            path = export_to_docx(content, title, filename)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            filename += '.pdf'
            path = export_to_pdf(content, title, filename)
            mimetype = 'application/pdf'

        return send_file(path, mimetype=mimetype,
                        as_attachment=True,
                        download_name=filename)
    except ImportError as e:
        return jsonify({"error": f"Library not available: {str(e)}. Install with pip install python-docx reportlab"}), 500
    except Exception as e:
        return jsonify({"error": f"Error exporting: {str(e)}"}), 500


@app.route('/api/status', methods=['GET'])
def status():
    """Checks if Ollama is available."""
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=5)
        models = [m['name'] for m in resp.json().get('models', [])]
        return jsonify({
            "ollama": True,
            "models": models,
            "active_model": OLLAMA_MODEL
        })
    except:
        return jsonify({
            "ollama": False,
            "models": [],
            "active_model": OLLAMA_MODEL
        })


if __name__ == '__main__':
    print("=" * 50)
    print("  NeuroAgent Backend")
    print("  http://localhost:5000")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)